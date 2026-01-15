import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration
MD_FILE_PATH = "Final-Project/SDD/NavSight_SDD_Final.md"
# New filename to avoid locks
DOCX_OUTPUT_PATH = "Final-Project/SDD/NavSight_SDD_Final_v4.docx" 
ASSETS_DIR = "Final-Project/SDD/assets"

# Mapping diagrams to their filenames
DIAGRAM_MAP = {
    "architecture_diagram.png": "System Architecture Diagram",
    "class_diagram.png": "Class Diagram",
    "sequence_diagram.png": "Sequence Diagram",
    "activity_diagram.png": "Activity Diagram"
}

def create_sdd_docx():
    # 1. Read Markdown Content
    with open(MD_FILE_PATH, "r", encoding="utf-8") as f:
        md_content = f.read()

    # 2. Initialize Document
    doc = Document()
    
    # Title Page Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # 3. Parse and Write Content
    lines = md_content.split('\n')
    
    # Simple state machine for parsing
    in_code_block = False
    code_block_type = ""
    diagram_index = 0
    diagram_files = list(DIAGRAM_MAP.keys())

    for line in lines:
        line = line.strip()
        
        # Handle Headers
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
            
        # Handle Images / Mermaid Blocks
        elif line.startswith("```mermaid"):
            in_code_block = True
            code_block_type = "mermaid"
            continue
        elif line.startswith("```"):
            if in_code_block and code_block_type == "mermaid":
                # End of mermaid block - Insert Image
                if diagram_index < len(diagram_files):
                    img_filename = diagram_files[diagram_index]
                    img_path = os.path.join(ASSETS_DIR, img_filename)
                    
                    if os.path.exists(img_path):
                        # Add caption
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(f"Figure {diagram_index + 1}: {DIAGRAM_MAP[img_filename]}")
                        run.italic = True
                        
                        # Add image
                        try:
                            doc.add_picture(img_path, width=Inches(6.0))
                            last_paragraph = doc.paragraphs[-1] 
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            doc.add_paragraph(f"[Error inserting image: {e}]")
                    else:
                        doc.add_paragraph(f"[Image not found: {img_filename}]")
                    
                    diagram_index += 1
            
            in_code_block = not in_code_block
            code_block_type = ""
            continue
            
        # Skip content inside mermaid blocks
        elif in_code_block and code_block_type == "mermaid":
            continue
            
        # Handle Code Blocks (non-mermaid)
        elif in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            p.runs[0].font.name = 'Consolas'
            p.runs[0].font.size = Pt(9)
            continue

        # Handle Lists
        elif line.startswith("* ") or line.startswith("- "):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            # Remove "1. " and add as list item
            text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(text, style='List Number')
            
        # Handle Normal Text
        else:
            if line:
                # Handle Bold formatting (**text**)
                parts = re.split(r'(\*\*.*\*\*)', line)
                p = doc.add_paragraph()
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

    # 4. Save
    doc.save(DOCX_OUTPUT_PATH)
    print(f"Successfully generated: {DOCX_OUTPUT_PATH}")

if __name__ == "__main__":
    create_sdd_docx()